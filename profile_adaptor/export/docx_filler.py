"""Fill DOCX template or source resume with adapted content (layout-preserving)."""

from __future__ import annotations

import copy
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from profile_adaptor.models import AdaptedResume, FillReport
from profile_adaptor.parse.resume_parser import SECTION_ALIASES, _norm


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph text while keeping the first run's formatting when possible."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _heading_key(text: str) -> Optional[str]:
    n = _norm(text).rstrip(":")
    for key, aliases in SECTION_ALIASES.items():
        if n in aliases:
            return key
    return None


def _adapted_blocks(adapted: AdaptedResume) -> Dict[str, List[str]]:
    skills_line = ", ".join(adapted.skills) if adapted.skills else ""
    exp_lines: List[str] = []
    for e in adapted.experience:
        header = " — ".join(x for x in [e.employer, e.title, e.time_range] if x)
        if header:
            exp_lines.append(header)
        for b in e.bullets:
            exp_lines.append(f"• {b}")
        if exp_lines and exp_lines[-1] != "":
            exp_lines.append("")
    edu_lines: List[str] = []
    for e in adapted.education:
        header = " — ".join(x for x in [e.school, e.degree, e.time_range] if x)
        if header:
            edu_lines.append(header)
        if e.details:
            edu_lines.append(e.details)
        if edu_lines and edu_lines[-1] != "":
            edu_lines.append("")
    return {
        "summary": [adapted.summary] if adapted.summary else [],
        "skills": [skills_line] if skills_line else [],
        "experience": [l for l in exp_lines if l != ""],
        "education": [l for l in edu_lines if l != ""],
    }


def _clone_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """Insert a new paragraph after `paragraph`, cloning its style/pPr when possible."""
    new_p = copy.deepcopy(paragraph._p)
    # Clear runs in the clone
    for child in list(new_p):
        if child.tag == qn("w:r"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def _write_section_content(
    doc: Document,
    paragraphs: List[Paragraph],
    start: int,
    end: int,
    content: List[str],
) -> Tuple[List[Paragraph], int]:
    """
    Fill body slots between start and end; insert cloned paragraphs when content overflows.
    Returns updated paragraphs list and number of lines written.
    """
    body_slots = list(range(start + 1, end))
    written = 0

    if not content:
        for slot in body_slots:
            _set_paragraph_text(paragraphs[slot], "")
        return paragraphs, 0

    if not body_slots:
        # Insert after heading
        anchor = paragraphs[start]
        last = anchor
        for line in content:
            last = _clone_paragraph_after(last, line)
            written += 1
        # Refresh paragraph list
        return list(doc.paragraphs), written

    for j, slot in enumerate(body_slots):
        if j < len(content):
            _set_paragraph_text(paragraphs[slot], content[j])
            written += 1
        else:
            _set_paragraph_text(paragraphs[slot], "")

    if len(content) > len(body_slots):
        anchor = paragraphs[body_slots[-1]]
        last = anchor
        for line in content[len(body_slots) :]:
            last = _clone_paragraph_after(last, line)
            written += 1
        paragraphs = list(doc.paragraphs)

    return paragraphs, written


def _fill_contact_block(
    paragraphs: List[Paragraph],
    indices: Dict[str, int],
    contact: str,
) -> bool:
    if not contact or not paragraphs:
        return False
    first_section_idx = min(indices.values()) if indices else len(paragraphs)
    # Prefer paragraphs before first known section heading
    contact_lines = [l for l in contact.splitlines() if l.strip()]
    if not contact_lines:
        return False
    slots = list(range(0, first_section_idx))
    if not slots:
        _set_paragraph_text(paragraphs[0], contact_lines[0])
        return True
    for j, slot in enumerate(slots):
        if j < len(contact_lines):
            _set_paragraph_text(paragraphs[slot], contact_lines[j])
        # leave remaining pre-heading paras untouched if fewer contact lines
    if len(contact_lines) > len(slots):
        # join overflow into last contact slot
        _set_paragraph_text(
            paragraphs[slots[-1]],
            "\n".join(contact_lines[len(slots) - 1 :]),
        )
    return True


def fill_docx(
    template_path: Path,
    adapted: AdaptedResume,
    output_path: Path,
    contact_override: Optional[str] = None,
) -> Tuple[Path, FillReport]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    doc = Document(str(output_path))
    paragraphs = list(doc.paragraphs)
    blocks = _adapted_blocks(adapted)
    report = FillReport()

    indices: Dict[str, int] = {}
    for i, p in enumerate(paragraphs):
        key = _heading_key(p.text or "")
        if key and key not in indices:
            indices[key] = i

    contact_text = (contact_override or adapted.contact or "").strip()
    if contact_text:
        if _fill_contact_block(paragraphs, indices, contact_text):
            report.sections_filled.append("contact")
        else:
            report.sections_missing.append("contact")
            report.notes.append("Could not place contact block.")
    else:
        report.sections_missing.append("contact")

    if not indices:
        report.degraded = True
        report.notes.append(
            "No recognizable section headings; wrote structured content without collapsing into one paragraph."
        )
        # Clear body but rebuild with proper headings (do NOT dump into para 0)
        for p in paragraphs:
            _set_paragraph_text(p, "")
        if contact_text:
            if paragraphs:
                _set_paragraph_text(paragraphs[0], contact_text.split("\n")[0][:120])
                for line in contact_text.splitlines()[1:]:
                    doc.add_paragraph(line)
            else:
                doc.add_paragraph(contact_text)
        for title, key in (
            ("Summary", "summary"),
            ("Skills", "skills"),
            ("Experience", "experience"),
            ("Education", "education"),
        ):
            doc.add_heading(title, level=1)
            lines = blocks.get(key) or []
            if not lines:
                doc.add_paragraph("")
                report.sections_missing.append(key)
            else:
                for line in lines:
                    doc.add_paragraph(line)
                report.sections_filled.append(key)
        doc.save(str(output_path))
        return output_path, report

    # Recompute indices after contact edits
    paragraphs = list(doc.paragraphs)
    indices = {}
    for i, p in enumerate(paragraphs):
        key = _heading_key(p.text or "")
        if key and key not in indices:
            indices[key] = i

    keys_in_order = sorted(indices.items(), key=lambda kv: kv[1])
    for idx, (key, start) in enumerate(keys_in_order):
        # Re-resolve end against current paragraph list
        paragraphs = list(doc.paragraphs)
        # refresh start index
        start = next(
            (i for i, p in enumerate(paragraphs) if _heading_key(p.text or "") == key),
            start,
        )
        later = [
            i
            for i, p in enumerate(paragraphs)
            if i > start and _heading_key(p.text or "") in {k for k, _ in keys_in_order}
        ]
        end = later[0] if later else len(paragraphs)
        content = blocks.get(key) or []
        paragraphs, written = _write_section_content(doc, paragraphs, start, end, content)
        if written:
            report.sections_filled.append(key)
        else:
            report.sections_missing.append(key)

    expected = {"summary", "skills", "experience", "education"}
    for key in expected - set(report.sections_filled):
        if key not in report.sections_missing:
            report.sections_missing.append(key)

    doc.save(str(output_path))
    return output_path, report


def create_docx_from_adapted(adapted: AdaptedResume, output_path: Path) -> Tuple[Path, FillReport]:
    """Create a clean DOCX when no usable template exists."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    report = FillReport(degraded=True, notes=["Created new DOCX (no template fill)."])
    if adapted.contact:
        doc.add_heading(adapted.contact.split("\n")[0][:80], level=0)
        for line in adapted.contact.splitlines()[1:]:
            doc.add_paragraph(line)
        report.sections_filled.append("contact")
    if adapted.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(adapted.summary)
        report.sections_filled.append("summary")
    if adapted.skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(adapted.skills))
        report.sections_filled.append("skills")
    if adapted.experience:
        doc.add_heading("Experience", level=1)
        for e in adapted.experience:
            header = " — ".join(x for x in [e.employer, e.title, e.time_range] if x)
            doc.add_paragraph(header)
            for b in e.bullets:
                doc.add_paragraph(b, style="List Bullet")
        report.sections_filled.append("experience")
    if adapted.education:
        doc.add_heading("Education", level=1)
        for e in adapted.education:
            header = " — ".join(x for x in [e.school, e.degree, e.time_range] if x)
            doc.add_paragraph(header)
            if e.details:
                doc.add_paragraph(e.details)
        report.sections_filled.append("education")
    if adapted.extras:
        doc.add_heading("Additional", level=1)
        doc.add_paragraph(adapted.extras)
    doc.save(str(output_path))
    return output_path, report
