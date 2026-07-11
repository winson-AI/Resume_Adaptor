"""Parse DOCX/PDF resumes into structured ResumeDocument (fast, quiet PDF path)."""

from __future__ import annotations

import logging
import re
import warnings
from pathlib import Path
from typing import List, Tuple

from docx import Document
from pypdf import PdfReader

from profile_adaptor.models import EducationEntry, ExperienceEntry, ResumeDocument

# pdfplumber/pdfminer emit noisy FontBBox warnings on many Chinese/resume PDFs
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfminer.pdfinterp").setLevel(logging.ERROR)
logging.getLogger("pdfminer.pdffont").setLevel(logging.ERROR)

SECTION_ALIASES = {
    "summary": [
        "summary",
        "profile",
        "objective",
        "about",
        "professional summary",
        "个人简介",
        "自我评价",
        "概述",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "expertise",
        "技能",
        "专业技能",
        "技术栈",
    ],
    "experience": [
        "experience",
        "work experience",
        "employment",
        "professional experience",
        "工作经历",
        "工作经验",
        "项目经历",
    ],
    "education": [
        "education",
        "academic",
        "学历",
        "教育背景",
        "教育经历",
    ],
}

_HEADING_LOOKUP = {alias: key for key, aliases in SECTION_ALIASES.items() for alias in aliases}
_TIME_RANGE_RE = re.compile(
    r"((?:19|20)\d{2}\s*[-–—to至]+\s*(?:(?:19|20)\d{2}|present|now|至今|今))",
    re.I,
)
_BULLET_RE = re.compile(r"^[\-•·*]\s*")
_EXP_SPLIT_RE = re.compile(r"\n(?=[A-Z\u4e00-\u9fff].{2,80})")


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower())


def _is_heading(line: str) -> Tuple[bool, str]:
    n = _norm(line).rstrip(":")
    if not n or len(n) > 60:
        return False, ""
    key = _HEADING_LOOKUP.get(n)
    return (True, key) if key else (False, "")


def _split_sections(text: str) -> dict:
    sections = {
        "contact": "",
        "summary": "",
        "skills": "",
        "experience": "",
        "education": "",
        "extras": "",
    }
    current = "contact"
    buckets = {k: [] for k in sections}
    for raw in text.splitlines():
        line = raw.rstrip()
        is_h, key = _is_heading(line)
        if is_h:
            current = key
            continue
        buckets[current].append(line)
    for k, lines in buckets.items():
        sections[k] = "\n".join(lines).strip()
    return sections


def _parse_skills(block: str) -> List[str]:
    if not block:
        return []
    parts = re.split(r"[,;|•·\n]+", block)
    return [p.strip() for p in parts if p.strip() and len(p.strip()) < 80]


def _parse_experience(block: str) -> List[ExperienceEntry]:
    if not block:
        return []
    entries: List[ExperienceEntry] = []
    chunks = _EXP_SPLIT_RE.split(block)
    if len(chunks) == 1:
        chunks = [c for c in re.split(r"\n{2,}", block) if c.strip()]
    for chunk in chunks:
        lines = [l.strip() for l in chunk.splitlines() if l.strip()]
        if not lines:
            continue
        header = lines[0]
        time_m = _TIME_RANGE_RE.search(chunk)
        bullets = [_BULLET_RE.sub("", l) for l in lines[1:] if l]
        employer, title = header, ""
        if " - " in header or " – " in header or "—" in header:
            parts = re.split(r"\s[-–—]\s", header, maxsplit=1)
            if len(parts) == 2:
                employer, title = parts[0].strip(), parts[1].strip()
        elif "|" in header:
            parts = [p.strip() for p in header.split("|", 1)]
            employer, title = parts[0], parts[1] if len(parts) > 1 else ""
        entries.append(
            ExperienceEntry(
                employer=employer,
                title=title,
                time_range=time_m.group(1).strip() if time_m else "",
                bullets=bullets or ([lines[0]] if len(lines) == 1 else []),
            )
        )
    return entries


def _parse_education(block: str) -> List[EducationEntry]:
    if not block:
        return []
    entries: List[EducationEntry] = []
    chunks = [c for c in re.split(r"\n{2,}", block) if c.strip()]
    if len(chunks) == 1:
        chunks = [c for c in block.splitlines() if c.strip()]
        if len(chunks) > 1:
            entries.append(
                EducationEntry(
                    school=chunks[0],
                    degree=chunks[1] if len(chunks) > 1 else "",
                    time_range="",
                    details="\n".join(chunks[2:]),
                )
            )
            return entries
    for chunk in chunks:
        lines = [l.strip() for l in chunk.splitlines() if l.strip()]
        if not lines:
            continue
        time_m = _TIME_RANGE_RE.search(chunk)
        entries.append(
            EducationEntry(
                school=lines[0],
                degree=lines[1] if len(lines) > 1 else "",
                time_range=time_m.group(1).strip() if time_m else "",
                details="\n".join(lines[2:]),
            )
        )
    return entries


def _extract_pdf_with_pypdf(path: Path) -> str:
    reader = PdfReader(str(path), strict=False)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts).strip()


def _extract_pdf_with_pdfplumber(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        return ""
    parts = []
    # Suppress pdfminer FontBBox / font descriptor noise
    with warnings.catch_warnings():
        warnings.filterwarnings(
            "ignore",
            message=r".*FontBBox.*",
        )
        warnings.filterwarnings(
            "ignore",
            message=r".*font descriptor.*",
            category=UserWarning,
        )
        try:
            with pdfplumber.open(str(path)) as pdf:
                # Resumes are usually short; cap pages for speed
                for page in pdf.pages[:12]:
                    try:
                        parts.append(page.extract_text() or "")
                    except Exception:
                        continue
        except Exception:
            return ""
    return "\n".join(parts).strip()


def _extract_pdf_text(path: Path) -> str:
    """Prefer pypdf (fast, quiet); fall back to pdfplumber only if text is thin."""
    text = _extract_pdf_with_pypdf(path)
    # Enough signal for sectioning? otherwise try pdfplumber once
    if len(text) >= 120 and text.count("\n") >= 3:
        return text
    alt = _extract_pdf_with_pdfplumber(path)
    if len(alt) > len(text):
        return alt
    return text


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    # Avoid building huge lists for empty runs
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def parse_resume(path: str) -> ResumeDocument:
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"Resume not found: {path}")
    suffix = p.suffix.lower()
    if suffix == ".docx":
        raw = _extract_docx_text(p)
        fmt = "docx"
    elif suffix == ".pdf":
        raw = _extract_pdf_text(p)
        fmt = "pdf"
    else:
        raise ValueError(f"Unsupported resume format: {suffix} (use .docx or .pdf)")

    if not raw.strip():
        raise ValueError(f"Could not extract text from resume: {path}")

    sections = _split_sections(raw)
    return ResumeDocument(
        contact=sections["contact"],
        summary=sections["summary"],
        skills=_parse_skills(sections["skills"]),
        experience=_parse_experience(sections["experience"]),
        education=_parse_education(sections["education"]),
        extras=sections["extras"],
        raw_text=raw,
        source_path=str(p.resolve()),
        source_format=fmt,
    )
